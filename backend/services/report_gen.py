from docx import Document
from io import BytesIO

def generate_docx_report(analysis: dict, stats: dict):
    doc = Document()
    doc.add_heading('Academic NLP Analysis Report', 0)

    # Section 1: Thematic
    doc.add_heading('1. Thematic Analysis', level=1)
    doc.add_paragraph(analysis['thematic'])

    # Section 2: Writing Style
    doc.add_heading('2. Lexical & Writing Style', level=1)
    doc.add_paragraph(analysis['lexical'])
    
    # Add a small table for stats
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Grade Level'
    row_cells[1].text = str(stats['lexical']['grade_level'])

    # Save to memory (BytesIO) so we can send it via FastAPI
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream