from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from pydantic import BaseModel, Field
from typing import Optional, TypedDict
from database import check_db_connection
from datetime import datetime, UTC
from database import db
from models import User, UserLogin, AnalysisRequest
from auth_utils import hash_password, create_access_token, verify_password
from services.nlp_engine import TextProcessor
from services.agents import agent_executor
from fastapi.responses import StreamingResponse
from services.report_gen import generate_docx_report
from bson import ObjectId
from docx import Document
import io
import datetime as dt

from fastapi.middleware.cors import CORSMiddleware
app = FastAPI(title="Academic NLP Analyzer")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173", "http://localhost:3000"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- SCHEMAS ---
class AnalysisState(TypedDict):
    raw_data: dict
    text: str
    style: str
    audience: str
    intent: str
    vision_summary: str
    technical_gaps: list
    thematic_paragraph: str
    sentiment_paragraph: str
    lexical_paragraph: str
    ai_signal_paragraph: str
    master_guidance: str

# --- UTILS ---
def validate_input(text: str):
    words = text.split()
    if len(words) < 30:
        return "Input is too short for a meaningful writing analysis (Minimum 30 words)."
    special_chars = sum(1 for char in text if char in "{}[]()<>|_")
    if special_chars > (len(words) * 0.6):
        return "Input detected as raw data or code. Please provide prose for analysis."
    return None

async def run_analysis(text: str, style: str, audience: str, intent: str, title: str = "Untitled", save: bool = False, user_id: str = None):
    processor = TextProcessor(text)
    raw_data = {
        "lexical": processor.get_lexical_stats(),
        "themes": processor.get_thematic_keywords(),
        "sentiment": processor.get_sentiment_metrics(),
        "readability": processor.get_readability_metrics(),
        "ai_signal": processor.get_ai_signal(),
    }
    initial_state = {
        "raw_data": raw_data, 
        "text": text,
        "style": style or "General",      
        "audience": audience or "General", 
        "intent": intent or "Informative",
        "vision_summary": "",
        "technical_gaps": []
    }
    final_state = await agent_executor.ainvoke(initial_state)
    analysis_results = {
        "vision": final_state.get("vision_summary"),
        "thematic": final_state.get("thematic_paragraph"),
        "sentiment": final_state.get("sentiment_paragraph"),
        "lexical": final_state.get("lexical_paragraph"),
        "ai_signal": final_state.get("ai_signal_paragraph"),
        "master_action_plan": final_state.get("master_guidance")
    }
    report_id = None
    if save:
        if not user_id:
            raise HTTPException(status_code=400, detail="user_id is required")
        report_doc = {
            "user_id": user_id,
            "title": title or "Untitled Analysis",
            "content_preview": text[:100] + "...",
            "analysis": analysis_results,
            "raw_stats": raw_data,
            "created_at": dt.datetime.now(dt.UTC)
        }
        result = await db.reports.insert_one(report_doc)
        report_id = str(result.inserted_id)
    return {
        "status": "success",
        "report_id": report_id,
        "analysis": analysis_results,
        "raw_stats": raw_data
    }

# --- ROUTES ---
@app.on_event("startup")
async def startup_event():
    await check_db_connection()

@app.get("/")
async def health_check():
    return {"status": "online", "message": "Academic NLP Backend is running"}

@app.post("/auth/signup")
async def signup(user_data: User):
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    user_dict = user_data.dict()
    user_dict["password"] = hash_password(user_data.password)
    new_user = await db.users.insert_one(user_dict)
    token = create_access_token(data={"sub": user_data.email})
    return {"message": "User created", "access_token": token, "user_id": str(new_user.inserted_id), "token_type": "bearer"}

@app.post("/auth/login")
async def login(user_credentials: UserLogin):
    user = await db.users.find_one({"email": user_credentials.email})
    if not user or not verify_password(user_credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    access_token = create_access_token(data={"sub": user["email"]})
    return {"access_token": access_token, "user_id": str(user["_id"]), "token_type": "bearer"}

@app.post("/analyze/text")
async def analyze_text(request: AnalysisRequest, save: bool = False, user_id: str = None):
    error_message = validate_input(request.text)
    if error_message:
        raise HTTPException(status_code=400, detail=error_message)
    return await run_analysis(request.text, request.style, request.audience, request.intent, request.title, save, user_id)

@app.post("/extract/file")
async def extract_file(file: UploadFile = File(...)):
    MAX_FILE_SIZE = 5 * 1024 * 1024
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")
    
    text = ""
    file_extension = file.filename.split(".")[-1].lower()
    try:
        if file_extension == "txt":
            text = content.decode("utf-8")
        elif file_extension == "docx":
            doc = Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload .txt or .docx")
        
        return {"text": text, "filename": file.filename}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text: {str(e)}")

@app.post("/analyze/file")
async def analyze_file(
    file: UploadFile = File(...), 
    style: str = Form("General"), 
    audience: str = Form("General"), 
    intent: str = Form("Informative"),
    save: bool = Form(False),
    user_id: str = Form(None)
):
    try:
        MAX_FILE_SIZE = 5 * 1024 * 1024
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")
        
        text = ""
        file_extension = file.filename.split(".")[-1].lower()
        if file_extension == "txt":
            text = content.decode("utf-8")
        elif file_extension == "docx":
            doc = Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload .txt or .docx")
            
        error_message = validate_input(text)
        if error_message:
            raise HTTPException(status_code=400, detail=error_message)
            
        # 4. Run Analysis
        result = await run_analysis(text, style, audience, intent, file.filename, save, user_id)
        result["text"] = text  # Include extracted text in response
        return result
    except HTTPException:
        raise
    except Exception as e:
        print(f"ANALYZE_FILE ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal Server Error during file analysis: {str(e)}")

@app.get("/users/{user_id}/reports")
async def get_user_reports(user_id: str):
    reports = []
    async for report in db.reports.find({"user_id": user_id}).sort("created_at", -1):
        reports.append({
            "id": str(report["_id"]),
            "title": report.get("title", "Untitled Analysis"),
            "content_preview": report.get("content_preview", ""),
            "created_at": report["created_at"]
        })
    return reports

@app.get("/reports/{report_id}")
async def get_report_detail(report_id: str):
    report = await db.reports.find_one({"_id": ObjectId(report_id)})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return {
        "id": str(report["_id"]),
        "title": report.get("title"),
        "analysis": report.get("analysis"),
        "raw_stats": report.get("raw_stats"),
        "created_at": report.get("created_at")
    }

@app.get("/reports/{report_id}/download")
async def download_report(report_id: str):
    report = await db.reports.find_one({"_id": ObjectId(report_id)})
    if not report: raise HTTPException(status_code=404, detail="Report not found")
    file = generate_docx_report(report['analysis'], report['raw_stats'])
    return StreamingResponse(file, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=Analysis_Report.docx"})

@app.get("/reports/{report_id}/export/text")
async def download_report_text(report_id: str):
    report = await db.reports.find_one({"_id": ObjectId(report_id)})
    if not report: raise HTTPException(status_code=404, detail="Report not found")
    
    analysis = report['analysis']
    
    text_content = f"GRIFFITH WRITING INSIGHTS - ARCHIVAL REPORT\n"
    text_content += f"==========================================\n\n"
    text_content += f"TITLE: {report.get('title', 'Untitled')}\n"
    text_content += f"DATE: {report.get('created_at')}\n"
    text_content += f"------------------------------------------\n\n"
    
    for section, content in analysis.items():
        text_content += f"[{section.upper()}]\n"
        if isinstance(content, dict):
            text_content += f"STATUS: {content.get('verdict', 'N/A')}\n"
            text_content += f"KEY INSIGHT: {content.get('headline', 'N/A')}\n"
            signals = ", ".join(content.get('signals', []))
            text_content += f"SIGNALS: {signals}\n"
            text_content += f"\nDETAILED ANALYSIS:\n{content.get('analysis', '')}\n"
        else:
            text_content += f"{content}\n"
        text_content += f"\n{'='*40}\n\n"
            
    return StreamingResponse(
        io.BytesIO(text_content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename=Griffith_Analysis_{report_id}.txt"}
    )